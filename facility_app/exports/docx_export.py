"""
DOCX export: fills the approved Facility_Assessment_Snapshot.docx
template in-place, preserving its exact logo, fonts, borders, and
column widths - rather than rebuilding the document from scratch.

The template's table cells contain placeholder text ("<Text>",
"{Address}", etc.) in row order matching mapping.build_snapshot_fields()
exactly, so we fill by row index rather than by searching for specific
placeholder strings (more robust to minor placeholder wording changes).
"""

import io
from pathlib import Path
from docx import Document

# Anchor to this module's location, not the process's CWD - Streamlit Cloud
# doesn't guarantee CWD == the repo/app folder.
BASE_DIR = Path(__file__).resolve().parent.parent
TEMPLATE_PATH = BASE_DIR / "assets" / "snapshot_template.docx"


def _set_cell_text(cell, text: str):
    """
    Replace a table cell's text while preserving the existing run
    formatting (italic placeholder style) of its first run.
    """
    if not cell.paragraphs:
        cell.text = text
        return
    paragraph = cell.paragraphs[0]
    if paragraph.runs:
        # Reuse the first run's formatting, clear the rest
        paragraph.runs[0].text = text
        for extra_run in paragraph.runs[1:]:
            extra_run.text = ""
    else:
        paragraph.add_run(text)
    # Clear any additional paragraphs in the cell beyond the first
    for p in cell.paragraphs[1:]:
        for r in p.runs:
            r.text = ""


def build_snapshot_docx(snapshot: dict, ccn: str) -> bytes:
    """
    snapshot: the dict returned by mapping.build_snapshot_fields()
    ccn: included for interface parity with build_snapshot_pdf (the
         hyperlink isn't required for the docx export by the case study,
         but we add it as a courtesy at the end of the document).
    Returns DOCX file bytes.
    """
    doc = Document(TEMPLATE_PATH)

    # Fill {STATE} header paragraph - it's the 2nd non-empty paragraph
    state_value = snapshot.get("state") or ""
    for p in doc.paragraphs:
        if p.text.strip() == "{STATE}":
            if p.runs:
                p.runs[0].text = state_value
                for extra in p.runs[1:]:
                    extra.text = ""
            break

    # Fill the 24-row table, by row index - matches mapping.py row order
    table = doc.tables[0]
    rows = snapshot["rows"]
    for i, (_, value) in enumerate(rows):
        if i >= len(table.rows):
            break
        _set_cell_text(table.rows[i].cells[1], str(value))

    # Courtesy: add the Medicare hyperlink as a plain paragraph at the end
    # (not a clickable hyperlink in docx - python-docx hyperlinks require
    # extra XML relationship work; the PDF export carries the clickable
    # link, which is the case study's actual hard requirement)
    state = (snapshot.get("state") or "").strip().upper()
    url = f"https://www.medicare.gov/care-compare/details/nursing-home/{ccn.strip()}/view-all?state={state}"
    doc.add_paragraph()
    doc.add_paragraph(f"Medicare Care Compare profile: {url}")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
